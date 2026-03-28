"""Document ingestion: parse PDF, EPUB, DOCX, and TXT into structured sections."""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class Paragraph:
    """A single paragraph of text."""
    text: str
    page_number: int | None = None


@dataclass
class Section:
    """A structural section (chapter, heading, etc.)."""
    heading: str
    level: int  # 1 = top-level chapter, 2 = sub-section, etc.
    paragraphs: list[Paragraph] = field(default_factory=list)
    children: list["Section"] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """The result of parsing a document."""
    title: str | None
    author: str | None
    page_count: int | None
    sections: list[Section]
    raw_text: str  # Fallback: full concatenated text


def parse_document(file_path: str) -> ParsedDocument:
    """Parse a document file into structured sections.

    Dispatches to format-specific parsers based on file extension.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    parsers = {
        ".pdf": _parse_pdf,
        ".epub": _parse_epub,
        ".docx": _parse_docx,
        ".txt": _parse_txt,
    }

    parser = parsers.get(ext)
    if not parser:
        raise ValueError(f"Unsupported file format: {ext}")

    logger.info("Parsing %s (%s)", path.name, ext)
    return parser(path)


# ---------------------------------------------------------------------------
# PDF Parser
# ---------------------------------------------------------------------------

def _parse_pdf(path: Path) -> ParsedDocument:
    """Extract text and structure from a PDF.

    Uses pdfplumber for text extraction. Falls back to OCR for scanned pages.
    """
    import pdfplumber

    sections: list[Section] = []
    all_text_parts: list[str] = []
    current_section = Section(heading="Document", level=1)

    with pdfplumber.open(path) as pdf:
        page_count = len(pdf.pages)
        title = pdf.metadata.get("Title")
        author = pdf.metadata.get("Author")

        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""

            # If no text extracted, attempt OCR
            if not text.strip():
                text = _ocr_page(page, page_num)

            if not text.strip():
                continue

            all_text_parts.append(text)

            # Simple heuristic: detect headings by line analysis
            lines = text.split("\n")
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    continue

                if _looks_like_heading(stripped):
                    # Save previous section
                    if current_section.paragraphs:
                        sections.append(current_section)
                    current_section = Section(
                        heading=stripped,
                        level=_guess_heading_level(stripped),
                    )
                else:
                    current_section.paragraphs.append(
                        Paragraph(text=stripped, page_number=page_num)
                    )

    # Don't forget the last section
    if current_section.paragraphs:
        sections.append(current_section)

    return ParsedDocument(
        title=title,
        author=author,
        page_count=page_count,
        sections=sections or [Section(heading="Full Text", level=1, paragraphs=[
            Paragraph(text=t) for t in all_text_parts if t.strip()
        ])],
        raw_text="\n\n".join(all_text_parts),
    )


def _ocr_page(page, page_num: int) -> str:
    """OCR a single PDF page using pytesseract."""
    try:
        from PIL import Image
        import pytesseract

        from app.config import settings

        if settings.tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

        # Convert page to image
        img = page.to_image(resolution=300).original
        text = pytesseract.image_to_string(img)
        logger.info("OCR extracted %d chars from page %d", len(text), page_num)
        return text
    except ImportError:
        logger.warning(
            "pytesseract not installed — skipping OCR for page %d", page_num
        )
        return ""
    except Exception as e:
        logger.warning("OCR failed on page %d: %s", page_num, e)
        return ""


# ---------------------------------------------------------------------------
# EPUB Parser
# ---------------------------------------------------------------------------

def _parse_epub(path: Path) -> ParsedDocument:
    """Extract text and structure from an EPUB."""
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    book = epub.read_epub(str(path))

    title = book.get_metadata("DC", "title")
    title = title[0][0] if title else None
    author = book.get_metadata("DC", "creator")
    author = author[0][0] if author else None

    sections: list[Section] = []
    all_text_parts: list[str] = []

    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "lxml")

        # Extract headings and paragraphs
        current_section = Section(
            heading=item.get_name(),
            level=1,
        )

        for element in soup.find_all(["h1", "h2", "h3", "h4", "p"]):
            text = element.get_text(strip=True)
            if not text:
                continue

            if element.name.startswith("h"):
                if current_section.paragraphs:
                    sections.append(current_section)
                level = int(element.name[1])
                current_section = Section(heading=text, level=level)
            else:
                current_section.paragraphs.append(Paragraph(text=text))
                all_text_parts.append(text)

        if current_section.paragraphs:
            sections.append(current_section)

    return ParsedDocument(
        title=title,
        author=author,
        page_count=None,
        sections=sections,
        raw_text="\n\n".join(all_text_parts),
    )


# ---------------------------------------------------------------------------
# DOCX Parser
# ---------------------------------------------------------------------------

def _parse_docx(path: Path) -> ParsedDocument:
    """Extract text and structure from a DOCX."""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    core = doc.core_properties

    sections: list[Section] = []
    all_text_parts: list[str] = []
    current_section = Section(heading="Document", level=1)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        all_text_parts.append(text)

        # Detect headings by style
        style_name = (para.style.name or "").lower()
        if "heading" in style_name:
            if current_section.paragraphs:
                sections.append(current_section)
            # Extract level from style name (e.g., "Heading 2" → 2)
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            current_section = Section(heading=text, level=level)
        else:
            current_section.paragraphs.append(Paragraph(text=text))

    if current_section.paragraphs:
        sections.append(current_section)

    return ParsedDocument(
        title=core.title,
        author=core.author,
        page_count=None,
        sections=sections,
        raw_text="\n\n".join(all_text_parts),
    )


# ---------------------------------------------------------------------------
# TXT Parser
# ---------------------------------------------------------------------------

def _parse_txt(path: Path) -> ParsedDocument:
    """Parse a plain text file into sections based on blank-line separators."""
    text = path.read_text(encoding="utf-8", errors="replace")
    paragraphs = [
        Paragraph(text=p.strip())
        for p in text.split("\n\n")
        if p.strip()
    ]

    return ParsedDocument(
        title=path.stem,
        author=None,
        page_count=None,
        sections=[Section(heading="Full Text", level=1, paragraphs=paragraphs)],
        raw_text=text,
    )


# ---------------------------------------------------------------------------
# Heading detection helpers
# ---------------------------------------------------------------------------

_HEADING_KEYWORDS = {"chapter", "part", "section", "appendix", "introduction",
                     "conclusion", "preface", "foreword", "epilogue", "prologue"}


def _looks_like_heading(line: str) -> bool:
    """Heuristic: is this line likely a section heading?"""
    # Short lines in all caps
    if len(line) < 80 and line.isupper():
        return True
    # Starts with "Chapter", "Part", numbered like "1.", "1.2"
    lower = line.lower().split()
    if lower and lower[0].rstrip(".") in _HEADING_KEYWORDS:
        return True
    if lower and lower[0].replace(".", "").isdigit() and len(line) < 100:
        return True
    return False


def _guess_heading_level(line: str) -> int:
    """Guess the structural level of a heading."""
    lower = line.lower()
    if lower.startswith(("part ", "book ")):
        return 1
    if lower.startswith("chapter"):
        return 2
    if lower.startswith("section"):
        return 3
    return 2
